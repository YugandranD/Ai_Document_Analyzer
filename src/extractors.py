"""
Text extraction utilities for PDF, DOCX, and image files.
Handles base64 decoding, OCR via Tesseract, and layout-aware extraction.
"""

import base64
import io
import os
import platform
import logging

from PyPDF2 import PdfReader
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)


def _configure_tesseract():
    """Configure Tesseract path based on OS."""
    from src.config import TESSERACT_PATH

    if platform.system() == "Windows" and os.path.exists(TESSERACT_PATH):
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH


def decode_base64(file_base64: str) -> bytes:
    """Decode a base64-encoded string to raw bytes."""
    # Handle data URIs (e.g. "data:application/pdf;base64,...")
    if "," in file_base64 and file_base64.startswith("data:"):
        file_base64 = file_base64.split(",", 1)[1]
    return base64.b64decode(file_base64)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file, falling back to OCR for scanned pages."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages_text.append(text.strip())
        else:
            # Scanned page — attempt OCR (if images are extractable)
            logger.info(f"Page {i+1}: no text layer, attempting OCR fallback")
            try:
                ocr_text = _ocr_pdf_page_fallback(page)
                if ocr_text:
                    pages_text.append(ocr_text)
            except Exception as e:
                logger.warning(f"OCR fallback failed for page {i+1}: {e}")

    return "\n\n".join(pages_text)


def _ocr_pdf_page_fallback(page) -> str:
    """Try to OCR images embedded in a PDF page."""
    import pytesseract
    _configure_tesseract()

    texts = []
    if hasattr(page, "images"):
        for img_obj in page.images:
            try:
                image = Image.open(io.BytesIO(img_obj.data))
                text = pytesseract.image_to_string(image)
                if text.strip():
                    texts.append(text.strip())
            except Exception:
                continue
    return "\n".join(texts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file preserving paragraph structure."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def extract_text_from_image(file_bytes: bytes) -> str:
    """Extract text from an image using Tesseract OCR."""
    import pytesseract
    _configure_tesseract()

    image = Image.open(io.BytesIO(file_bytes))

    # Convert to RGB if necessary (handles RGBA, palette, etc.)
    if image.mode not in ("L", "RGB"):
        image = image.convert("RGB")

    text = pytesseract.image_to_string(image)
    return text.strip()


def extract_text(file_type: str, file_bytes: bytes) -> str:
    """
    Route extraction to the appropriate handler based on file type.

    Args:
        file_type: One of 'pdf', 'docx', 'image'
        file_bytes: Raw file bytes

    Returns:
        Extracted text content
    """
    file_type = file_type.lower().strip()

    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif file_type in ("image", "img", "png", "jpg", "jpeg", "tiff", "bmp"):
        return extract_text_from_image(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
